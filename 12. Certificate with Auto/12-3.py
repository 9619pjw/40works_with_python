import docx
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = docx.Document(r'12. Certificate with Auto\form.docx')

# 폰트 및 사이즈 설정
style = doc.styles['Normal']
style.font.name = '나눔고딕'
style._element.rPr.rFonts.set(qn('w:eastAsia'), '나눔고딕')
style.font.size = docx.shared.Pt(12)

# 문단 생성 및 글자 입력
para = doc.add_paragraph()
run = para.add_run('\n\n')
run = para.add_run('\t  제 2020-0001 호 \n')
run.font.name = '나눔고딕'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '나눔고딕')
run.font.size = docx.shared.Pt(20)

para = doc.add_paragraph()
run = para.add_run('\n\n')
run = para.add_run(' 수 료 증 \n')
run.font.name = '나눔고딕'
run.bold = True
run._element.rPr.rFonts.set(qn('w:eastAsia'), '나눔고딕')
run.font.size = docx.shared.Pt(40)
para.alignment = WD_ALIGN_PARAGRAPH.CENTER

para = doc.add_paragraph()
run = para.add_run('\n\n')
run = para.add_run('\t  성 명 : 장다인 \n')
run.font.name = '나눔고딕'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '나눔고딕')
run.font.size = docx.shared.Pt(20)
run = para.add_run('\t  생 년 월 일 : 2017.12.12  \n')
run.font.name = '나눔고딕'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '나눔고딕')
run.font.size = docx.shared.Pt(20)
run = para.add_run('\t  교 육 과 정 : 40 Works with Python  \n')
run.font.name = '나눔고딕'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '나눔고딕')
run.font.size = docx.shared.Pt(20)
run = para.add_run('\t  교 육 날 짜 : 2024.04.01 ~ 2024.08.01 \n')
run.font.name = '나눔고딕'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '나눔고딕')
run.font.size = docx.shared.Pt(20)

para = doc.add_paragraph()
run = para.add_run('\n\n')
run = para.add_run('\t  위 사람은 해당 교육과정을 \n')
run.font.name = '나눔고딕'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '나눔고딕')
run.font.size = docx.shared.Pt(20)
run = para.add_run('\t  이수하였으므로 이 증서를 수여합니다. \n')
run.font.name = '나눔고딕'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '나눔고딕')
run.font.size = docx.shared.Pt(20)

para = doc.add_paragraph()
run = para.add_run('\n\n\n')
run = para.add_run('2024.08.09 \n')
run.font.name = '나눔고딕'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '나눔고딕')
run.font.size = docx.shared.Pt(20)
para.alignment = WD_ALIGN_PARAGRAPH.CENTER

para = doc.add_paragraph()
run = para.add_run('\n\n\n')
run = para.add_run('@@교육기관장 \n')
run.font.name = '나눔고딕'
run.bold = True
run._element.rPr.rFonts.set(qn('w:eastAsia'), '나눔고딕')
run.font.size = docx.shared.Pt(20)
para.alignment = WD_ALIGN_PARAGRAPH.CENTER

# 파일 저장
doc.save(r'12. Certificate with Auto\form.docx')